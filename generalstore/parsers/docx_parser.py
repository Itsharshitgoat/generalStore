"""
DOCX parser using python-docx.

Chunks document content by heading hierarchy (Heading 1 / Heading 2).
If no headings are present, paragraphs are batched into groups of ~10.
Tables are serialized in markdown format and emitted as separate chunks.
"""

import logging
from pathlib import Path
from typing import List, Optional

from docx import Document

from generalstore.config import get_settings
from generalstore.parsers.base import BaseParser, DocumentChunk

logger = logging.getLogger(__name__)

_HEADING_STYLES = {"Heading 1", "Heading 2"}
_PARAGRAPHS_PER_BATCH = 10


class DOCXParser(BaseParser):
    """Parser for .docx files using python-docx."""

    supported_extensions: List[str] = [".docx"]

    def parse(self, filepath: Path) -> List[DocumentChunk]:
        """
        Parse a DOCX file and return a list of DocumentChunks.

        Content is chunked by heading hierarchy (H1/H2).  If the
        document contains no headings, paragraphs are batched in
        groups of ~10.  Tables are serialized as markdown and added
        as separate chunks.

        Args:
            filepath: Absolute path to the .docx file.

        Returns:
            List of DocumentChunk objects.

        Raises:
            FileNotFoundError: If the file does not exist.
            RuntimeError: If DOCX parsing fails.
        """
        filepath = Path(filepath).resolve()
        if not filepath.exists():
            raise FileNotFoundError(f"DOCX file not found: {filepath}")

        settings = get_settings()
        subject = self._derive_subject(filepath)
        chunks: List[DocumentChunk] = []
        chunk_index = 0

        try:
            doc = Document(str(filepath))
        except Exception as exc:
            raise RuntimeError(
                f"Failed to open DOCX '{filepath.name}': {exc}"
            ) from exc

        # ----- Paragraph chunks (heading-based or batched) -----
        has_headings = any(
            p.style and p.style.name in _HEADING_STYLES
            for p in doc.paragraphs
        )

        if has_headings:
            chunk_index = self._chunk_by_headings(
                doc, filepath, settings, subject, chunks, chunk_index
            )
        else:
            chunk_index = self._chunk_by_batch(
                doc, filepath, settings, subject, chunks, chunk_index
            )

        # ----- Table chunks -----
        for table_idx, table in enumerate(doc.tables):
            try:
                md = self._table_to_markdown(table)
                cleaned = self._clean_text(md)
                if len(cleaned) < settings.min_chunk_chars:
                    continue
                chunks.append(
                    DocumentChunk(
                        text=cleaned,
                        source_file=str(filepath),
                        file_type=".docx",
                        chunk_index=chunk_index,
                        heading=f"Table {table_idx + 1}",
                        subject=subject,
                    )
                )
                chunk_index += 1
            except Exception as exc:
                logger.warning(
                    "Failed to parse table %d in '%s': %s",
                    table_idx + 1,
                    filepath.name,
                    exc,
                )

        logger.info(
            "Parsed DOCX '%s': %d chunks", filepath.name, len(chunks)
        )
        return chunks

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _chunk_by_headings(
        self,
        doc: Document,
        filepath: Path,
        settings,
        subject: Optional[str],
        chunks: List[DocumentChunk],
        chunk_index: int,
    ) -> int:
        """Split paragraphs into chunks at every Heading 1 / Heading 2."""
        current_heading: Optional[str] = None
        buffer: List[str] = []

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""

            if style_name in _HEADING_STYLES:
                # Flush the previous section
                chunk_index = self._flush_buffer(
                    buffer,
                    current_heading,
                    filepath,
                    settings,
                    subject,
                    chunks,
                    chunk_index,
                )
                buffer = []
                current_heading = self._clean_text(para.text)
            else:
                text = para.text.strip()
                if text:
                    buffer.append(text)

        # Flush remaining
        chunk_index = self._flush_buffer(
            buffer,
            current_heading,
            filepath,
            settings,
            subject,
            chunks,
            chunk_index,
        )
        return chunk_index

    def _chunk_by_batch(
        self,
        doc: Document,
        filepath: Path,
        settings,
        subject: Optional[str],
        chunks: List[DocumentChunk],
        chunk_index: int,
    ) -> int:
        """Batch paragraphs into groups of ~10 when no headings exist."""
        buffer: List[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            buffer.append(text)

            if len(buffer) >= _PARAGRAPHS_PER_BATCH:
                chunk_index = self._flush_buffer(
                    buffer,
                    None,
                    filepath,
                    settings,
                    subject,
                    chunks,
                    chunk_index,
                )
                buffer = []

        # Flush remaining
        chunk_index = self._flush_buffer(
            buffer,
            None,
            filepath,
            settings,
            subject,
            chunks,
            chunk_index,
        )
        return chunk_index

    def _flush_buffer(
        self,
        buffer: List[str],
        heading: Optional[str],
        filepath: Path,
        settings,
        subject: Optional[str],
        chunks: List[DocumentChunk],
        chunk_index: int,
    ) -> int:
        """Join buffered paragraphs, clean, and append as a chunk."""
        if not buffer:
            return chunk_index

        combined = "\n".join(buffer)
        cleaned = self._clean_text(combined)

        if len(cleaned) < settings.min_chunk_chars:
            return chunk_index

        # Respect max_chunk_chars – split if necessary
        segments = self._split_text(cleaned, settings.max_chunk_chars)
        for seg in segments:
            if len(seg) < settings.min_chunk_chars:
                continue
            chunks.append(
                DocumentChunk(
                    text=seg,
                    source_file=str(filepath),
                    file_type=".docx",
                    chunk_index=chunk_index,
                    heading=heading,
                    subject=subject,
                )
            )
            chunk_index += 1
        return chunk_index

    @staticmethod
    def _table_to_markdown(table) -> str:
        """Serialize a python-docx Table object into markdown table format."""
        rows = table.rows
        if not rows:
            return ""

        lines: List[str] = []

        # Header row
        header_cells = [cell.text.strip() for cell in rows[0].cells]
        lines.append("| " + " | ".join(header_cells) + " |")
        lines.append("| " + " | ".join("---" for _ in header_cells) + " |")

        # Data rows
        for row in rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")

        return "\n".join(lines)

    @staticmethod
    def _split_text(text: str, max_chars: int) -> List[str]:
        """Split text into segments of at most *max_chars*, breaking at boundaries."""
        if len(text) <= max_chars:
            return [text]

        parts: List[str] = []
        while text:
            if len(text) <= max_chars:
                parts.append(text)
                break

            split_pos = text.rfind(". ", 0, max_chars)
            if split_pos == -1 or split_pos < max_chars // 2:
                split_pos = text.rfind(" ", 0, max_chars)
            if split_pos == -1:
                split_pos = max_chars

            parts.append(text[: split_pos + 1].strip())
            text = text[split_pos + 1 :].strip()

        return parts
